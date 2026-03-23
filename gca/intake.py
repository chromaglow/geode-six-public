"""
Geode Six — GCA File Intake
Handles file upload, AI-powered naming, date resolution, and commit flow.

Two-step flow:
  1. POST /gca/upload → AI suggests filename, returns confirmation payload
  2. POST /gca/confirm → User confirms/edits, file is written to GCA

v2: Two-tier folder structure (Projects / Operations), dynamic codes from codes.json.
"""

import json
import logging
import os
import re
import shutil
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

import httpx
from dotenv import load_dotenv
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel

from gca.codes import (
    TYPE_CODES,
    build_naming_prompt,
    load_codes,
    tier_for_code,
    valid_code,
    all_folder_codes,
    GCA_ROOT,
)

load_dotenv()

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://localhost:11434")
LOG_PATH = os.getenv("LOG_PATH", "/mnt/nvme/logs")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".xlsx", ".jpg", ".jpeg", ".png"}

# AI naming model
NAMING_MODEL = "geode-llama31"

# Temp storage for pending uploads (keyed by temp_id)
_pending_uploads: dict = {}

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
os.makedirs(LOG_PATH, exist_ok=True)
intake_logger = logging.getLogger("geode.intake")
intake_logger.setLevel(logging.INFO)
log_file = os.path.join(LOG_PATH, "intake.log")
_fh = logging.FileHandler(log_file)
_fh.setFormatter(logging.Formatter("%(asctime)s %(message)s"))
intake_logger.addHandler(_fh)

# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------
router = APIRouter()

# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------


class UploadResponse(BaseModel):
    temp_id: str
    original_filename: str
    suggested_filename: str
    tier: str
    project: str
    type: str
    description: str
    date: str
    date_estimated: bool
    version: str
    duplicate_warning: Optional[str] = None


class ConfirmRequest(BaseModel):
    temp_id: str
    tier: str
    project: str
    type: str
    description: str
    date: str
    version: str
    ready_to_share: bool = False


class ConfirmResponse(BaseModel):
    assigned_filename: str
    stored_path: str
    indexed: bool


# ---------------------------------------------------------------------------
# Date resolution
# ---------------------------------------------------------------------------


def resolve_date(
    user_note: Optional[str],
    file_metadata_date: Optional[str],
    original_filename: str,
    file_ctime: Optional[float],
) -> tuple[str, bool]:
    """
    Resolve date using priority order:
      1. User-provided date in note
      2. Document metadata date
      3. Date from filename
      4. File system ctime
      5. Today's date (flagged as estimated)

    Returns (YYYYMMDD string, date_estimated bool).
    """
    # 1. User note
    if user_note:
        parsed = _parse_date_from_text(user_note)
        if parsed:
            return parsed, False

    # 2. File metadata
    if file_metadata_date:
        parsed = _parse_date_from_text(file_metadata_date)
        if parsed:
            return parsed, False

    # 3. Filename
    parsed = _parse_date_from_text(original_filename)
    if parsed:
        return parsed, False

    # 4. ctime
    if file_ctime:
        dt = datetime.fromtimestamp(file_ctime)
        return dt.strftime("%Y%m%d"), False

    # 5. Today (estimated)
    return datetime.now().strftime("%Y%m%d"), True


def _parse_date_from_text(text: str) -> Optional[str]:
    """Try to extract a date from text in various formats."""
    if not text:
        return None

    # YYYYMMDD
    m = re.search(r"(20\d{2})(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])", text)
    if m:
        return m.group(0)

    # YYYY-MM-DD or YYYY/MM/DD
    m = re.search(r"(20\d{2})[-/](0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])", text)
    if m:
        return f"{m.group(1)}{m.group(2)}{m.group(3)}"

    # MM-DD-YY or MM/DD/YY or M/D/YY
    m = re.search(r"(\d{1,2})[-/](\d{1,2})[-/](\d{2})\b", text)
    if m:
        month = int(m.group(1))
        day = int(m.group(2))
        year = int(m.group(3)) + 2000
        if 1 <= month <= 12 and 1 <= day <= 31:
            return f"{year}{month:02d}{day:02d}"

    # MonDD or Mon DD (e.g., Mar13, Mar 13)
    months = {
        "jan": "01", "feb": "02", "mar": "03", "apr": "04",
        "may": "05", "jun": "06", "jul": "07", "aug": "08",
        "sep": "09", "oct": "10", "nov": "11", "dec": "12",
    }
    m = re.search(r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s*(\d{1,2})", text, re.I)
    if m:
        month_str = months.get(m.group(1).lower())
        day = int(m.group(2))
        if month_str and 1 <= day <= 31:
            year = datetime.now().year
            return f"{year}{month_str}{day:02d}"

    return None


# ---------------------------------------------------------------------------
# Filename construction
# ---------------------------------------------------------------------------


def build_filename(
    project: str, type_code: str, description: str, date: str, version: str, ext: str
) -> str:
    """Build a GCA-compliant filename. No spaces allowed."""
    # Remove any spaces from description
    description = description.replace(" ", "")
    name = f"{project}_{type_code}_{description}_{date}_v{version}{ext}"
    assert " " not in name, f"Filename contains spaces: {name}"
    return name


def check_duplicate(tier: str, project: str, type_code: str, description: str) -> Optional[str]:
    """Check if a file with same PROJECT + TYPE + Description already exists."""
    project_dir = os.path.join(GCA_ROOT, tier, project)
    if not os.path.isdir(project_dir):
        return None

    prefix = f"{project}_{type_code}_{description}_"
    for fname in os.listdir(project_dir):
        if fname.startswith(prefix):
            # Extract version from existing file
            m = re.search(r"_v(\d+\.\d+)", fname)
            existing_version = m.group(1) if m else "unknown"
            return (
                f"A v{existing_version} of this file already exists. "
                f"This will be saved as the next version. Continue?"
            )
    return None


def next_version(tier: str, project: str, type_code: str, description: str, base_version: str) -> str:
    """Find the next available version number."""
    project_dir = os.path.join(GCA_ROOT, tier, project)
    if not os.path.isdir(project_dir):
        return base_version

    prefix = f"{project}_{type_code}_{description}_"
    max_major = 0
    max_minor = 0

    for fname in os.listdir(project_dir):
        if fname.startswith(prefix):
            m = re.search(r"_v(\d+)\.(\d+)", fname)
            if m:
                major = int(m.group(1))
                minor = int(m.group(2))
                if major > max_major or (major == max_major and minor > max_minor):
                    max_major = major
                    max_minor = minor

    if max_major == 0 and max_minor == 0:
        return base_version

    # Increment minor version
    return f"{max_major}.{max_minor + 1}"


# ---------------------------------------------------------------------------
# File metadata extraction
# ---------------------------------------------------------------------------


def extract_metadata_date(filepath: str, ext: str) -> Optional[str]:
    """Try to extract a date from file metadata."""
    try:
        if ext == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            info = reader.metadata
            if info and info.creation_date:
                return info.creation_date.strftime("%Y%m%d")

        elif ext == ".docx":
            from docx import Document
            doc = Document(filepath)
            props = doc.core_properties
            if props.created:
                return props.created.strftime("%Y%m%d")

        elif ext in (".jpg", ".jpeg", ".png"):
            from PIL import Image
            from PIL.ExifTags import TAGS
            img = Image.open(filepath)
            exif = img._getexif()
            if exif:
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id)
                    if tag == "DateTimeOriginal":
                        # Format: "YYYY:MM:DD HH:MM:SS"
                        return value.replace(":", "").split(" ")[0]
    except Exception:
        pass
    return None


def extract_text_preview(filepath: str, ext: str, max_chars: int = 2000) -> str:
    """Extract text preview from a file for AI naming."""
    try:
        if ext == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            text = ""
            for page in reader.pages[:3]:
                text += page.extract_text() or ""
                if len(text) > max_chars:
                    break
            return text[:max_chars]

        elif ext == ".docx":
            from docx import Document
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs[:30])
            return text[:max_chars]

        elif ext in (".txt", ".md"):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(max_chars)

        elif ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(filepath, read_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(max_row=20, values_only=True):
                rows.append(" | ".join(str(c) for c in row if c))
            return "\n".join(rows)[:max_chars]

        elif ext in (".jpg", ".jpeg", ".png"):
            return f"[Image file: {os.path.basename(filepath)}]"

    except Exception as e:
        return f"[Could not extract text: {e}]"
    return ""


# ---------------------------------------------------------------------------
# AI naming
# ---------------------------------------------------------------------------


async def ai_suggest_name(text_preview: str, original_filename: str, note: Optional[str]) -> dict:
    """Ask Llama 3.1 to suggest tier, code, type, and description."""
    user_msg = f"Original filename: {original_filename}\n"
    if note:
        user_msg += f"User note: {note}\n"
    user_msg += f"\nDocument content preview:\n{text_preview[:1500]}"

    naming_prompt = build_naming_prompt()

    payload = {
        "model": NAMING_MODEL,
        "prompt": f"{naming_prompt}\n\n{user_msg}",
        "stream": False,
    }

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(f"{OLLAMA_HOST}/api/generate", json=payload)
            resp.raise_for_status()
            result = resp.json()
            response_text = result.get("response", "")

            # Parse JSON from response (handle potential markdown wrapping)
            json_match = re.search(r'\{[^}]+\}', response_text)
            if json_match:
                suggestion = json.loads(json_match.group())

                # Extract fields
                tier = suggestion.get("tier", "Projects")
                code = suggestion.get("code", suggestion.get("project", "GEO")).upper()
                type_code = suggestion.get("type", "OPS").upper()
                description = suggestion.get("description", "Untitled")

                # Validate tier
                if tier not in ("Projects", "Operations"):
                    tier = "Projects"

                # Validate code against codes.json
                if not valid_code(code):
                    code = "GEO"
                    tier = "Projects"
                else:
                    # Ensure tier matches the code
                    actual_tier = tier_for_code(code)
                    if actual_tier:
                        tier = actual_tier

                if type_code not in TYPE_CODES:
                    type_code = "OPS"

                # Clean description (CamelCase, no spaces)
                description = description.replace(" ", "")

                return {"tier": tier, "code": code, "type": type_code, "description": description}

    except Exception as e:
        intake_logger.error(f"AI naming failed: {e}")

    # Fallback
    return {"tier": "Projects", "code": "GEO", "type": "OPS", "description": "UntitledDocument"}


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.post("/gca/upload", response_model=UploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    note: Optional[str] = Form(None),
    ready_to_share: bool = Form(False),
):
    """Upload a file. Returns AI-suggested naming for user review."""
    # Validate file type
    original_filename = file.filename or "unknown"
    ext = os.path.splitext(original_filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=(
                "This file type is not supported. "
                "Accepted types: PDF, DOCX, TXT, MD, XLSX, JPG, PNG"
            ),
        )

    # Save to temp location
    temp_id = str(uuid.uuid4())
    temp_dir = os.path.join(tempfile.gettempdir(), "geode_uploads")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, f"{temp_id}{ext}")

    with open(temp_path, "wb") as f:
        content = await file.read()
        f.write(content)

    intake_logger.info(f"Upload received: {original_filename} → temp_id={temp_id}")

    # Extract text preview for AI naming
    text_preview = extract_text_preview(temp_path, ext)

    # Extract metadata date
    metadata_date = extract_metadata_date(temp_path, ext)

    # Get file ctime
    file_ctime = os.path.getctime(temp_path)

    # AI naming
    suggestion = await ai_suggest_name(text_preview, original_filename, note)
    tier = suggestion["tier"]
    project = suggestion["code"]
    type_code = suggestion["type"]
    description = suggestion["description"]

    # Resolve date
    date, date_estimated = resolve_date(note, metadata_date, original_filename, file_ctime)

    # Version
    base_version = "1.0" if ready_to_share else "0.1"
    version = next_version(tier, project, type_code, description, base_version)

    # Build suggested filename
    suggested_filename = build_filename(project, type_code, description, date, version, ext)

    # Check for duplicates
    duplicate_warning = check_duplicate(tier, project, type_code, description)

    # Store pending upload
    _pending_uploads[temp_id] = {
        "temp_path": temp_path,
        "ext": ext,
        "original_filename": original_filename,
    }

    return UploadResponse(
        temp_id=temp_id,
        original_filename=original_filename,
        suggested_filename=suggested_filename,
        tier=tier,
        project=project,
        type=type_code,
        description=description,
        date=date,
        date_estimated=date_estimated,
        version=version,
        duplicate_warning=duplicate_warning,
    )


@router.post("/gca/confirm", response_model=ConfirmResponse)
async def confirm_upload(req: ConfirmRequest):
    """User confirms or edits naming fields. File is committed to GCA."""
    pending = _pending_uploads.get(req.temp_id)
    if not pending:
        raise HTTPException(status_code=404, detail="Upload not found or expired.")

    # Validate codes
    if not valid_code(req.project):
        raise HTTPException(status_code=400, detail=f"Invalid project code: {req.project}")
    if req.type not in TYPE_CODES:
        raise HTTPException(status_code=400, detail=f"Invalid type code: {req.type}")

    # Resolve tier (trust the request but validate against codes.json)
    tier = req.tier
    actual_tier = tier_for_code(req.project)
    if actual_tier:
        tier = actual_tier

    # Version override for ready_to_share
    version = req.version
    if req.ready_to_share and version.startswith("0."):
        version = "1.0"

    # Build final filename
    ext = pending["ext"]
    final_filename = build_filename(req.project, req.type, req.description, req.date, version, ext)

    # Ensure tier/project directory exists
    project_dir = os.path.join(GCA_ROOT, tier, req.project)
    os.makedirs(project_dir, exist_ok=True)

    # Copy file to final location
    final_path = os.path.join(project_dir, final_filename)

    # Never overwrite — increment version if file exists
    while os.path.exists(final_path):
        version = next_version(tier, req.project, req.type, req.description, version)
        final_filename = build_filename(req.project, req.type, req.description, req.date, version, ext)
        final_path = os.path.join(project_dir, final_filename)

    shutil.move(pending["temp_path"], final_path)

    # Remove from pending
    del _pending_uploads[req.temp_id]

    intake_logger.info(
        f"File committed: {pending['original_filename']} → {final_filename} "
        f"at {final_path}"
    )

    # Trigger Chroma embedding
    indexed = False
    try:
        from gca.embed import embed_file
        await embed_file(
            final_path,
            tier=tier,
            project=req.project,
            type_code=req.type,
            description=req.description,
            date=req.date,
            version=version,
        )
        indexed = True
    except ImportError:
        intake_logger.info("Embedding module not yet available — skipping indexing")
    except Exception as e:
        intake_logger.error(f"Embedding failed for {final_filename}: {e}")

    return ConfirmResponse(
        assigned_filename=final_filename,
        stored_path=final_path,
        indexed=indexed,
    )
